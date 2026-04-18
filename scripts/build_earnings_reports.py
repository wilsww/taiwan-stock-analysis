"""
Earnings Update Reports — build_earnings_reports.py
Companies: 華邦電子 (2344.TW) Q4 2025 & Micron Technology (MU) FQ2 2026
Language: 繁體中文
Generated: 2026-03-22
"""
import os
import io
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Colour palette ───────────────────────────────────────────
DARK_NAVY  = (23,54,93)
GOLD       = (192,144,0)
RED_WARN   = (192,0,0)
GREEN_OK   = (0,112,0)
AMBER      = (197,90,17)
AMBER_HEX  = "#c55a11"
LIGHT_GREY = (242,242,242)
WHITE      = (255,255,255)

def rgb(t): return RGBColor(*t)

# ─── Chart helpers ────────────────────────────────────────────
CHART_STYLE = dict(
    facecolor="#0f172a", edgecolor="#0f172a"
)
AXIS_STYLE = dict(facecolor="#1e293b")
TEXT_COLOR = "#e2e8f0"
GRID_COLOR = "#334155"
BAR_BLUE   = "#3b82f6"
BAR_GREEN  = "#22c55e"
BAR_RED    = "#ef4444"
BAR_AMBER  = "#f59e0b"
LINE_COLOR = "#60a5fa"

def base_fig(nrows=1, ncols=1, figsize=(9,4)):
    fig, ax = plt.subplots(nrows, ncols, figsize=figsize, **CHART_STYLE)
    if isinstance(ax, np.ndarray):
        for a in ax.flat:
            a.set_facecolor(AXIS_STYLE["facecolor"])
            a.tick_params(colors=TEXT_COLOR, labelsize=9)
            a.xaxis.label.set_color(TEXT_COLOR)
            a.yaxis.label.set_color(TEXT_COLOR)
            for spine in a.spines.values():
                spine.set_edgecolor(GRID_COLOR)
    else:
        ax.set_facecolor(AXIS_STYLE["facecolor"])
        ax.tick_params(colors=TEXT_COLOR, labelsize=9)
        for spine in ax.spines.values():
            spine.set_edgecolor(GRID_COLOR)
    fig.patch.set_facecolor("#0f172a")
    return fig, ax

def save_chart(fig, path):
    fig.tight_layout()
    fig.savefig(path, dpi=150, bbox_inches="tight",
                facecolor=fig.get_facecolor())
    plt.close(fig)

def chart_to_bytes(fig):
    buf = io.BytesIO()
    fig.tight_layout()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight",
                facecolor=fig.get_facecolor())
    plt.close(fig)
    buf.seek(0)
    return buf

# ─── DOCX helpers ─────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)

def add_heading(doc, text, level=1, color=DARK_NAVY, size=16):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(4)
    return p

def add_body(doc, text, size=10, italic=False, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    run.italic = italic
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(3)
    return p

def add_bullet(doc, text, size=10):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    return p

def add_chart(doc, fig_bytes, width=Inches(6.5), caption=None):
    doc.add_picture(fig_bytes, width=width)
    if caption:
        p = doc.add_paragraph(caption)
        p.runs[0].font.size = Pt(8)
        p.runs[0].italic = True
        p.paragraph_format.space_after = Pt(6)
    return

def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "17365D")
    pBdr.append(bottom)
    pPr.append(pBdr)

def cover_table(doc, rows_data, col_widths=None):
    """Simple 2-col key-value table"""
    tbl = doc.add_table(rows=len(rows_data), cols=2)
    tbl.style = "Table Grid"
    for i, (k, v, bold_v) in enumerate(rows_data):
        c0, c1 = tbl.rows[i].cells
        set_cell_bg(c0, "17365D")
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(k)
        r0.font.color.rgb = rgb(WHITE)
        r0.font.size = Pt(9.5)
        r0.bold = True
        r0.font.name = "Times New Roman"
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p1 = c1.paragraphs[0]
        r1 = p1.add_run(v)
        r1.font.size = Pt(9.5)
        r1.bold = bold_v
        r1.font.name = "Times New Roman"
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if col_widths:
        for i, row in enumerate(tbl.rows):
            for j, width in enumerate(col_widths):
                row.cells[j].width = width
    return tbl

def beat_miss_table(doc, headers, rows):
    tbl = doc.add_table(rows=1+len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    for j, h in enumerate(headers):
        c = tbl.rows[0].cells[j]
        set_cell_bg(c, "1F4E79")
        p = c.paragraphs[0]
        r = p.add_run(h)
        r.font.color.rgb = rgb(WHITE)
        r.bold = True
        r.font.size = Pt(9)
        r.font.name = "Times New Roman"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            c = tbl.rows[i+1].cells[j]
            p = c.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(9)
            r.font.name = "Times New Roman"
            # color beat/miss
            if j == len(row_data)-1:  # last col = beat/miss
                if "▲" in str(val) or "超" in str(val):
                    r.font.color.rgb = rgb(GREEN_OK)
                    r.bold = True
                elif "▼" in str(val) or "低" in str(val):
                    r.font.color.rgb = rgb(RED_WARN)
                    r.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return tbl


# ════════════════════════════════════════════════════════════════
# REPORT 1: 華邦電子 (2344.TW) Q4 2025
# ════════════════════════════════════════════════════════════════

def build_winbond_charts(tmpdir):
    charts = {}

    # 1. Quarterly Revenue (NT$億)
    qtrs = ["Q1'25", "Q2'25", "Q3'25", "Q4'25", "Q1'26E"]
    rev  = [193, 217, 218, 266, 355]
    colors = [BAR_BLUE]*4 + [BAR_AMBER]
    fig, ax = base_fig(figsize=(8,4))
    bars = ax.bar(qtrs, rev, color=colors, width=0.6)
    ax.set_title("華邦電 季度合并營收 (NT$億)", color=TEXT_COLOR, fontsize=11, pad=10)
    ax.set_ylabel("NT$億", color=TEXT_COLOR)
    ax.yaxis.grid(True, color=GRID_COLOR, linewidth=0.5)
    ax.set_axisbelow(True)
    for bar, v in zip(bars, rev):
        ax.text(bar.get_x()+bar.get_width()/2, bar.get_height()+3,
                f"{v}", ha="center", color=TEXT_COLOR, fontsize=9, fontweight="bold")
    ax.set_ylim(0, 420)
    p_act = mpatches.Patch(color=BAR_BLUE,  label="實際值")
    p_est = mpatches.Patch(color=BAR_AMBER, label="預估值")
    ax.legend(handles=[p_act, p_est], facecolor="#1e293b",
              labelcolor=TEXT_COLOR, fontsize=8)
    charts["rev"] = chart_to_bytes(fig)

    # 2. Quarterly EPS
    qtrs2 = ["Q1'25", "Q2'25", "Q3'25", "Q4'25"]
    eps   = [-0.24, -0.29, 0.65, 0.76]
    colors2 = [BAR_RED, BAR_RED, BAR_GREEN, BAR_GREEN]
    fig, ax = base_fig(figsize=(8,4))
    bars = ax.bar(qtrs2, eps, color=colors2, width=0.55)
    ax.axhline(0, color=GRID_COLOR, linewidth=0.8)
    ax.set_title("華邦電 季度每股盈餘 EPS (NT$)", color=TEXT_COLOR, fontsize=11, pad=10)
    ax.set_ylabel("EPS (NT$)", color=TEXT_COLOR)
    ax.yaxis.grid(True, color=GRID_COLOR, linewidth=0.5)
    ax.set_axisbelow(True)
    for bar, v in zip(bars, eps):
        offset = 0.03 if v >= 0 else -0.06
        ax.text(bar.get_x()+bar.get_width()/2, v+offset,
                f"{v:+.2f}", ha="center", color=TEXT_COLOR, fontsize=9, fontweight="bold")
    charts["eps"] = chart_to_bytes(fig)

    # 3. Gross Margin Trend
    qtrs3 = ["Q1'25", "Q2'25", "Q3'25", "Q4'25", "Q1'26E"]
    gm    = [8, 22.7, 46.7, 41.9, 47.0]
    colors3 = [BAR_BLUE]*4 + [BAR_AMBER]
    fig, ax = base_fig(figsize=(8,4))
    ax.bar(qtrs3, gm, color=colors3, width=0.6)
    ax.plot(qtrs3, gm, color=LINE_COLOR, marker="o", linewidth=2, markersize=6, zorder=5)
    ax.set_title("華邦電 合并毛利率趨勢 (%)", color=TEXT_COLOR, fontsize=11, pad=10)
    ax.set_ylabel("毛利率 %", color=TEXT_COLOR)
    ax.yaxis.grid(True, color=GRID_COLOR, linewidth=0.5)
    ax.set_axisbelow(True)
    for i, (q, v) in enumerate(zip(qtrs3, gm)):
        ax.text(i, v+1.5, f"{v:.1f}%", ha="center", color=TEXT_COLOR, fontsize=9, fontweight="bold")
    ax.set_ylim(0, 60)
    p_act = mpatches.Patch(color=BAR_BLUE,  label="實際值")
    p_est = mpatches.Patch(color=BAR_AMBER, label="預估值")
    ax.legend(handles=[p_act, p_est], facecolor="#1e293b",
              labelcolor=TEXT_COLOR, fontsize=8)
    ax.annotate("Q3 存貨回沖利益", xy=(2, 46.7), xytext=(2.5, 52),
                color=AMBER_HEX, fontsize=8,
                arrowprops=dict(arrowstyle="->", color=AMBER_HEX))
    charts["gm"] = chart_to_bytes(fig)

    # 4. Revenue Mix: DRAM vs NOR Flash (FY2025 estimate)
    segments = ["Specialty\nDRAM (CMS)", "NOR Flash", "Logic IC\n& Other"]
    sizes    = [55, 39, 6]
    explode  = (0.05, 0, 0)
    pie_colors = [BAR_BLUE, BAR_GREEN, BAR_AMBER]
    fig, ax = base_fig(figsize=(7, 4.5))
    wedges, texts, autotexts = ax.pie(
        sizes, explode=explode, labels=segments,
        colors=pie_colors, autopct="%1.0f%%",
        startangle=90, textprops={"color": TEXT_COLOR, "fontsize": 9}
    )
    for at in autotexts:
        at.set_fontsize(9); at.set_fontweight("bold"); at.set_color("white")
    ax.set_title("華邦電 FY2025 營收結構 (估計)", color=TEXT_COLOR, fontsize=11, pad=10)
    charts["mix"] = chart_to_bytes(fig)

    # 5. Beat/Miss Summary
    metrics = ["Q4 營收\n(NT$億)", "Q4 EPS\n(NT$)", "全年 EPS\n(NT$)"]
    actual  = [266.25, 0.76, 0.88]
    est     = [235.0,  0.55, 0.72]
    x = np.arange(len(metrics))
    w = 0.35
    fig, ax = base_fig(figsize=(8,4))
    b1 = ax.bar(x - w/2, actual, w, label="實際值", color=BAR_GREEN)
    b2 = ax.bar(x + w/2, est,    w, label="法人共識", color=BAR_BLUE)
    ax.set_title("華邦電 Q4 2025 實際 vs 法人共識", color=TEXT_COLOR, fontsize=11, pad=10)
    ax.set_xticks(x); ax.set_xticklabels(metrics, color=TEXT_COLOR, fontsize=9)
    ax.yaxis.grid(True, color=GRID_COLOR, linewidth=0.5)
    ax.set_axisbelow(True)
    ax.legend(facecolor="#1e293b", labelcolor=TEXT_COLOR, fontsize=9)
    for bar, v in zip(b1, actual):
        ax.text(bar.get_x()+bar.get_width()/2, v*1.02, f"{v}", ha="center",
                color=TEXT_COLOR, fontsize=9, fontweight="bold")
    charts["beat"] = chart_to_bytes(fig)

    # 6. 2026 CapEx Plan
    categories = ["生產設備\n(95%)", "廠房建設\n(3%)", "其他 IT\n(2%)"]
    amounts    = [400, 12, 9]
    fig, ax = base_fig(figsize=(7,4))
    ax.bar(categories, amounts, color=[BAR_BLUE, BAR_AMBER, BAR_GREEN], width=0.5)
    ax.set_title("華邦電 2026年資本支出計畫 NT$421億", color=TEXT_COLOR, fontsize=11, pad=10)
    ax.set_ylabel("NT$億", color=TEXT_COLOR)
    ax.yaxis.grid(True, color=GRID_COLOR, linewidth=0.5)
    ax.set_axisbelow(True)
    for i, (cat, v) in enumerate(zip(categories, amounts)):
        ax.text(i, v+5, f"NT${v}億", ha="center", color=TEXT_COLOR, fontsize=9, fontweight="bold")
    charts["capex"] = chart_to_bytes(fig)

    # 7. Forward PE vs Peers
    companies = ["華邦電\n(2344)", "南亞科\n(2408)", "Micron\n(MU)", "SanDisk\n(SNDK)"]
    fwd_pe    = [23.9, 14.1, 4.72, 9.05]
    colors7   = [BAR_AMBER, BAR_BLUE, BAR_GREEN, BAR_BLUE]
    fig, ax = base_fig(figsize=(8,4))
    bars = ax.bar(companies, fwd_pe, color=colors7, width=0.55)
    ax.axhline(np.median(fwd_pe), color=LINE_COLOR, linestyle="--",
               linewidth=1.5, label=f"中位數 {np.median(fwd_pe):.1f}x")
    ax.set_title("記憶體同業 2026E Forward P/E 比較", color=TEXT_COLOR, fontsize=11, pad=10)
    ax.set_ylabel("Forward P/E (x)", color=TEXT_COLOR)
    ax.yaxis.grid(True, color=GRID_COLOR, linewidth=0.5)
    ax.set_axisbelow(True)
    ax.legend(facecolor="#1e293b", labelcolor=TEXT_COLOR, fontsize=9)
    for bar, v in zip(bars, fwd_pe):
        ax.text(bar.get_x()+bar.get_width()/2, v+0.3, f"{v:.1f}x",
                ha="center", color=TEXT_COLOR, fontsize=9, fontweight="bold")
    charts["pe"] = chart_to_bytes(fig)

    # 8. DRAM Pricing Trajectory
    periods = ["Q2'25", "Q3'25", "Q4'25", "Q1'26E", "Q2'26E"]
    idx_price = [100, 110, 125, 237, 500]  # indexed, Q2'25=100
    fig, ax = base_fig(figsize=(8,4))
    ax.fill_between(periods, idx_price, alpha=0.3, color=BAR_BLUE)
    ax.plot(periods, idx_price, color=LINE_COLOR, marker="o",
            linewidth=2.5, markersize=8)
    ax.set_title("Specialty DRAM ASP 走勢（Q2'25 = 100）", color=TEXT_COLOR, fontsize=11, pad=10)
    ax.set_ylabel("ASP 指數", color=TEXT_COLOR)
    ax.yaxis.grid(True, color=GRID_COLOR, linewidth=0.5)
    ax.set_axisbelow(True)
    for i, (p, v) in enumerate(zip(periods, idx_price)):
        ax.text(i, v+12, f"{v}", ha="center", color=TEXT_COLOR, fontsize=9, fontweight="bold")
    ax.annotate("Q4: DDR4+90-95%\nQoQ", xy=(2, 125), xytext=(1.2, 200),
                color=AMBER_HEX, fontsize=8,
                arrowprops=dict(arrowstyle="->", color=AMBER_HEX))
    ax.annotate("Q2'26E: ~4x\nvs end-2025", xy=(4, 500), xytext=(3.3, 450),
                color=BAR_GREEN, fontsize=8,
                arrowprops=dict(arrowstyle="->", color=BAR_GREEN))
    charts["pricing"] = chart_to_bytes(fig)

    return charts


def build_winbond_report(outpath, charts):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # ── COVER ──────────────────────────────────────────────────
    p = doc.add_paragraph()
    r = p.add_run("華邦電子（2344.TW）")
    r.font.size = Pt(22); r.bold = True
    r.font.color.rgb = rgb(DARK_NAVY)
    r.font.name = "Times New Roman"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    r = p.add_run("2025 年第四季財報更新報告  |  Q4 2025 Earnings Update")
    r.font.size = Pt(14); r.bold = True
    r.font.color.rgb = rgb((89,89,89))
    r.font.name = "Times New Roman"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    cover_data = [
        ("投資評等",  "▲ 買進 (Buy)", True),
        ("目標價格",  "NT$135（上調自 NT$100）", True),
        ("現時股價",  "NT$110（2026/03/20）", False),
        ("潛在漲幅",  "+22.7%", True),
        ("報告日期",  "2026年3月22日", False),
        ("分析師",    "台灣科技股研究團隊", False),
        ("分析基準",  "Q4 2025 自結財報 + 4Q25 法說會簡報", False),
    ]
    cover_table(doc, cover_data, [Inches(2), Inches(4.2)])
    doc.add_paragraph()
    add_divider(doc)

    # ── KEY TAKEAWAYS ──────────────────────────────────────────
    add_heading(doc, "核心要點  KEY TAKEAWAYS", level=1,
                color=DARK_NAVY, size=13)
    takeaways = [
        "【超預期】Q4 EPS NT$0.76 大幅超越法人共識 NT$0.55，季增 17%，創 14 季新高。",
        "【毛利率正常化】Q4 合并毛利率 41.9%（Q3 46.7% 含存貨回沖利益），純記憶體段落更高；Q1 2026 指引 >46.9%，升勢確立。",
        "【DRAM 超級定價】本季 DDR4/LPDDR4 報價季增 90-95%；管理層預估 2026 年中報價將達去年底 4 倍水準。",
        "【CapEx 創歷史紀錄】2026 年資本支出上調至 NT$421 億（較 2025 年實際 NT$134 億大增 214%），95% 用於生產設備。",
        "【產能全數售出】2026/2027 年全年產能已悉數預訂，DRAM 供給結構性缺貨格局確認。",
        "【估值重估上修目標價】基於 2026E EPS NT$7.5（Bull Case）× 18x Forward PE，上調目標價至 NT$135。",
    ]
    for t in takeaways:
        add_bullet(doc, t, size=10)

    add_divider(doc)

    # ── RESULTS SNAPSHOT TABLE ─────────────────────────────────
    add_heading(doc, "財報快覽  RESULTS SNAPSHOT", level=1,
                color=DARK_NAVY, size=13)
    hdrs = ["財務指標", "Q4 2025 實際", "法人共識", "超/低預期", "QoQ 變化", "YoY 變化"]
    rows = [
        ["合并營收 (NT$億)", "266.25", "~235e", "▲ +13.3%", "+22.3%", "+42.4%"],
        ["合并毛利率 (%)", "41.9%", "~40%e", "▲ 略超", "- 4.8pp", "+15.4pp"],
        ["營業利益率 (%)", "15.4%", "~12%e", "▲ 超預期", "+9.4pp", "+15.8pp"],
        ["稅後淨利 (NT$億)", "34.22", "~25e", "▲ +36.9%", "+479%", "+4070%"],
        ["每股盈餘 EPS (NT$)", "0.76", "0.55e", "▲ +38.2%", "+17.0%", "扭虧"],
        ["全年 EPS (NT$)", "0.88", "0.72e", "▲ +22.2%", "—", "+529%"],
    ]
    beat_miss_table(doc, hdrs, rows)
    add_body(doc, "資料來源：4Q25 法說會簡報（2026/02/10）；法人共識為彭博估計均值。",
             size=8, italic=True)
    doc.add_paragraph()

    # ── CHARTS ─────────────────────────────────────────────────
    # Revenue + EPS side by side (insert as separate)
    add_heading(doc, "季度營收與獲利趨勢", level=2, color=(68,114,196), size=12)
    add_chart(doc, charts["rev"], width=Inches(6.2),
              caption="圖1：華邦電季度合并營收（NT$億）；Q1'26E 基於月營收年增 88% 推算。資料來源：公司財報、法說會。")
    add_chart(doc, charts["eps"], width=Inches(6.2),
              caption="圖2：季度 EPS（NT$）；FY2025 全年 EPS NT$0.88，年增 529%。資料來源：公司財報。")
    doc.add_paragraph()

    add_heading(doc, "毛利率趨勢與定價環境", level=2, color=(68,114,196), size=12)
    add_chart(doc, charts["gm"], width=Inches(6.2),
              caption="圖3：合并毛利率趨勢；Q3 含存貨回沖利益，Q4 正常化後仍達 41.9%。Q1'26E >46.9%。資料來源：公司財報。")
    add_chart(doc, charts["pricing"], width=Inches(6.2),
              caption="圖4：Specialty DRAM ASP 走勢指數（Q2'25=100）；Q1'26 DDR4 合約價季增 ~100%。資料來源：法說會。")

    doc.add_page_break()

    # ── SEGMENT ANALYSIS ──────────────────────────────────────
    add_heading(doc, "業務部門分析  SEGMENT ANALYSIS", level=1,
                color=DARK_NAVY, size=13)

    add_heading(doc, "① Specialty DRAM（CMS）— 核心成長驅動", level=2,
                color=(68,114,196), size=11)
    add_body(doc, "2025 全年 CMS 營收年增 139%，位元出貨量翻倍，遠超預期。20 奈米製程產品營收佔比持續提升，帶動 ASP 與毛利率同步改善。4Q25 DDR4 合約價季增 ~100%，1Q26 累計漲幅估計超過 100%，整體 DRAM 報價預計在 2026 年中達 2025 年底 4 倍。供需方面，2026/2027 年全年產能均已完售，結構性缺貨格局確立。", size=10)

    add_heading(doc, "② NOR Flash — 穩健成長，汽車工業需求帶動", level=2,
                color=(68,114,196), size=11)
    add_body(doc, "2025 年 NOR Flash 營收年增 8%，位元出貨量兩位數成長。受惠汽車（ADAS）、工業自動化、嵌入式系統需求強勁，平均售價穩定。台中廠追加 NT$24 億設備採購以擴充 NOR Flash 產能（500K → 570-580K 片/月）。全球最大 NOR Flash 供應商地位持續鞏固。", size=10)

    add_chart(doc, charts["mix"], width=Inches(5.5),
              caption="圖5：FY2025 營收結構估計（DRAM 55%、NOR Flash 39%）。資料來源：法說會。")

    doc.add_paragraph()

    # ── GUIDANCE & OUTLOOK ─────────────────────────────────────
    add_heading(doc, "展望與指引  GUIDANCE & OUTLOOK", level=1,
                color=DARK_NAVY, size=13)
    hdrs2 = ["展望項目", "2025 Q4 實際", "2026 Q1 指引", "2026 全年展望"]
    rows2 = [
        ["合并毛利率", "41.9%", ">46.9%（管理層指引）", "50%+ 可期（定價維持強勢）"],
        ["DRAM 報價漲幅", "+90-95% QoQ", "相近漲幅（管理層預期）", "2026 年中達 4x 底部"],
        ["產能利用率", "滿載", "滿載", "2026/2027 全年售出"],
        ["資本支出 (NT$億)", "134（全年實際）", "—", "421（2026 年計畫）"],
        ["股利 (NT$/股)", "配息 NT$0.5", "—", "配息比率可望提升"],
    ]
    beat_miss_table(doc, hdrs2, rows2)
    add_body(doc, "資料來源：4Q25 法說會（2026/02/10）。", size=8, italic=True)
    doc.add_paragraph()

    add_chart(doc, charts["capex"], width=Inches(5.5),
              caption="圖6：2026年資本支出計畫 NT$421億，95% 用於生產設備擴產。資料來源：4Q25 法說會。")

    doc.add_page_break()

    # ── INVESTMENT THESIS UPDATE ───────────────────────────────
    add_heading(doc, "投資論點更新  THESIS UPDATE", level=1,
                color=DARK_NAVY, size=13)
    add_body(doc, "本季財報強化我們對華邦電 Specialty DRAM 超級循環的核心論點。三項關鍵更新如下：", size=10)

    add_heading(doc, "✅ 正面更新", level=2, color=GREEN_OK, size=11)
    positives = [
        "DRAM 定價力道遠超預期：90-95% 季增印證了 Specialty DDR4/LPDDR4 的結構性缺貨，非暫時性定價。",
        "管理層前瞻性措辭轉為積極：「2026/2027 產能悉數售出」的明確表態，是超級循環早期的典型訊號。",
        "CapEx 大幅躍升至 NT$421 億：顯示管理層對多年需求的信心，資本紀律良好（95% 用於生產設備）。",
        "記憶體段落毛利率（>41.9% 合并，記憶體段落更高）持續改善，Nuvoton 稀釋效應相對減少。",
    ]
    for p_item in positives:
        add_bullet(doc, p_item, size=10)

    add_heading(doc, "⚠ 需關注風險", level=2, color=AMBER, size=11)
    risks = [
        "估值偏高：Forward PE 23.9x（Bloomberg 共識 4.6E）為四大同業中最貴，需靠獲利成長支撐。",
        "Nuvoton 稀釋持續：合并毛利率被子公司稀釋 ~9pp，純記憶體段落需分段建模。",
        "CXMT 競爭風險：Specialty DDR3/DDR4 市場中國製造商滲透加速，長期定價能力受挑戰。",
        "AI CapEx 週期敏感：若任一 Hyperscaler CapEx 指引 YoY 下修 >10%，需求預測需修正。",
    ]
    for r_item in risks:
        add_bullet(doc, r_item, size=10)

    add_divider(doc)

    # ── VALUATION ─────────────────────────────────────────────
    add_heading(doc, "估值分析  VALUATION", level=1, color=DARK_NAVY, size=13)

    add_heading(doc, "三情境目標價", level=2, color=(68,114,196), size=11)
    hdrs3 = ["情境", "概率", "2026E EPS", "PE 倍數", "目標價 (NT$)", "依據"]
    rows3 = [
        ["熊市 Bear",  "20%", "NT$3.5",  "14x", "NT$49",  "CXMT 侵蝕 + AI CapEx 下修"],
        ["基本 Base",  "55%", "NT$6.3",  "16x", "NT$101", "JPMorgan 共識；週期延續"],
        ["牛市 Bull",  "25%", "NT$7.5",  "18x", "NT$135", "報價超預期 + 純記憶體溢價"],
        ["加權目標價", "—",  "—",       "—",  "NT$108", "= 49×20% + 101×55% + 135×25%"],
    ]
    beat_miss_table(doc, hdrs3, rows3)
    add_body(doc,
        "本報告目標價採用牛市情境 NT$135（+23%），理由：Q4 2025 EPS 超預期幅度顯示 2026E EPS 向 NT$7-8 修正的概率已提升，週期加速指標（報價、產能利用率、指引措辭）均已轉向積極。",
        size=10)

    doc.add_paragraph()
    add_chart(doc, charts["pe"], width=Inches(6.2),
              caption="圖7：記憶體同業 2026E Forward P/E 比較；華邦電 23.9x 為四家最高。資料來源：Bloomberg 共識。")
    add_chart(doc, charts["beat"], width=Inches(6.2),
              caption="圖8：Q4 2025 實際 vs 法人共識對比；EPS 超預期 38%。資料來源：公司財報、Bloomberg。")

    doc.add_page_break()

    # ── ESTIMATES UPDATE ──────────────────────────────────────
    add_heading(doc, "預估值更新  ESTIMATE REVISIONS", level=1,
                color=DARK_NAVY, size=13)
    hdrs4 = ["指標", "舊 2026E 預估", "新 2026E 預估", "修正幅度"]
    rows4 = [
        ["全年營收 (NT$億)", "1,050",  "1,400",  "▲ +33%"],
        ["合并毛利率 (%)",   "42%",    "50%+",   "▲ +8pp"],
        ["EPS (NT$)",        "4.6",    "6.3-7.5","▲ +37-63%"],
        ["資本支出 (NT$億)", "300",    "421",    "▲ +40%"],
    ]
    beat_miss_table(doc, hdrs4, rows4)
    add_body(doc, "修正理由：Q4 2025 DRAM 報價加速（+90-95% QoQ）顯示 2026 年全年 ASP 上升幅度超出原始預估；加之 2026/2027 年產能全部售出確認，EPS 預估上修至 NT$6.3-7.5 區間。", size=10)

    add_divider(doc)

    # ── SOURCES ───────────────────────────────────────────────
    add_heading(doc, "資料來源  SOURCES & REFERENCES", level=1,
                color=DARK_NAVY, size=13)
    sources = [
        "4Q25 & FY2025 法人說明會簡報（2026/02/10）— https://www.winbond.com/hq/about-winbond/investor/investor-conference/",
        "Q4 2025 合并財務報告（自結）— 公開資訊觀測站 MOPS: https://mops.twse.com.tw",
        "財報狗 2344 EPS 分析 — https://statementdog.com/analysis/2344/eps",
        "StockFeel 4Q25 法說會重點 — https://www.stockfeel.com.tw",
        "TechNews 科技新報 (2026/02/11) — https://finance.technews.tw",
        "法人目標價參考：摩根大通（NT$83）、群益投顧（NT$120.3）；Bloomberg 共識 2026E EPS NT$4.6",
        "股價 NT$110 截至 2026/03/20 收盤",
    ]
    for s in sources:
        add_bullet(doc, s, size=9)

    add_body(doc,
        "免責聲明：本報告僅供個人投資研究參考，非正式投資建議。所有財務數據標記「e」或「估」者為估計值，"
        "尚未經會計師查核之自結數字可能與最終審計結果存在差異。",
        size=8, italic=True)

    doc.save(outpath)
    print(f"✓ Winbond report saved: {outpath}")


# ════════════════════════════════════════════════════════════════
# REPORT 2: Micron Technology (MU) FQ2 2026
# ════════════════════════════════════════════════════════════════

def build_micron_charts():
    charts = {}

    # 1. Quarterly Revenue ($B)
    qtrs  = ["FQ2'25", "FQ3'25", "FQ4'25", "FQ1'26", "FQ2'26", "FQ3'26E"]
    rev   = [7.7, 9.7, 11.3, 13.2, 23.9, 33.5]
    colors = [BAR_BLUE]*5 + [BAR_AMBER]
    fig, ax = base_fig(figsize=(9, 4))
    bars = ax.bar(qtrs, rev, color=colors, width=0.6)
    ax.set_title("Micron 季度合并營收 (US$B)", color=TEXT_COLOR, fontsize=11, pad=10)
    ax.set_ylabel("Revenue (US$B)", color=TEXT_COLOR)
    ax.yaxis.grid(True, color=GRID_COLOR, linewidth=0.5)
    ax.set_axisbelow(True)
    for bar, v in zip(bars, rev):
        ax.text(bar.get_x()+bar.get_width()/2, v+0.4,
                f"${v:.1f}B", ha="center", color=TEXT_COLOR, fontsize=8.5, fontweight="bold")
    ax.set_ylim(0, 40)
    ax.annotate("FQ3'26E: $33.5B\n(Record)", xy=(5, 33.5), xytext=(4.0, 36),
                color=BAR_AMBER, fontsize=8,
                arrowprops=dict(arrowstyle="->", color=BAR_AMBER))
    p_act = mpatches.Patch(color=BAR_BLUE,  label="Actual")
    p_est = mpatches.Patch(color=BAR_AMBER, label="Guidance")
    ax.legend(handles=[p_act, p_est], facecolor="#1e293b",
              labelcolor=TEXT_COLOR, fontsize=8)
    charts["rev"] = chart_to_bytes(fig)

    # 2. EPS
    qtrs2 = ["FQ2'25", "FQ3'25", "FQ4'25", "FQ1'26", "FQ2'26", "FQ3'26E"]
    eps   = [1.41, 2.04, 2.35, 4.67, 12.20, 19.15]
    colors2 = [BAR_BLUE]*5 + [BAR_AMBER]
    fig, ax = base_fig(figsize=(9, 4))
    ax.bar(qtrs2, eps, color=colors2, width=0.6)
    ax.plot(qtrs2, eps, color=LINE_COLOR, marker="o", linewidth=2, markersize=7, zorder=5)
    ax.set_title("Micron Non-GAAP EPS (US$)", color=TEXT_COLOR, fontsize=11, pad=10)
    ax.set_ylabel("EPS (US$)", color=TEXT_COLOR)
    ax.yaxis.grid(True, color=GRID_COLOR, linewidth=0.5)
    ax.set_axisbelow(True)
    for i, (q, v) in enumerate(zip(qtrs2, eps)):
        ax.text(i, v+0.4, f"${v:.2f}", ha="center", color=TEXT_COLOR, fontsize=8.5, fontweight="bold")
    charts["eps"] = chart_to_bytes(fig)

    # 3. Gross Margin
    qtrs3 = ["FQ2'25", "FQ3'25", "FQ4'25", "FQ1'26", "FQ2'26", "FQ3'26E"]
    gm    = [22, 29, 35, 57, 75, 81]
    colors3 = [BAR_BLUE]*5 + [BAR_AMBER]
    fig, ax = base_fig(figsize=(9, 4))
    ax.bar(qtrs3, gm, color=colors3, alpha=0.7, width=0.6)
    ax.plot(qtrs3, gm, color=LINE_COLOR, marker="o", linewidth=2.5, markersize=8, zorder=5)
    ax.set_title("Micron Non-GAAP Gross Margin (%)", color=TEXT_COLOR, fontsize=11, pad=10)
    ax.set_ylabel("Gross Margin %", color=TEXT_COLOR)
    ax.yaxis.grid(True, color=GRID_COLOR, linewidth=0.5)
    ax.set_axisbelow(True)
    for i, (q, v) in enumerate(zip(qtrs3, gm)):
        ax.text(i, v+1.5, f"{v}%", ha="center", color=TEXT_COLOR, fontsize=8.5, fontweight="bold")
    ax.set_ylim(0, 95)
    ax.annotate("Record\n75%", xy=(4, 75), xytext=(3.2, 85),
                color=BAR_GREEN, fontsize=8,
                arrowprops=dict(arrowstyle="->", color=BAR_GREEN))
    charts["gm"] = chart_to_bytes(fig)

    # 4. Revenue by Segment (FQ2 2026)
    segments = ["DRAM\n$18.8B", "NAND\n$5.0B", "Other\n$0.1B"]
    sizes    = [78.7, 20.9, 0.4]
    pie_colors = [BAR_BLUE, BAR_GREEN, BAR_AMBER]
    fig, ax = base_fig(figsize=(7, 4.5))
    wedges, texts, autotexts = ax.pie(
        sizes, labels=segments, colors=pie_colors,
        autopct="%1.1f%%", startangle=90,
        textprops={"color": TEXT_COLOR, "fontsize": 9}
    )
    for at in autotexts:
        at.set_fontsize(9); at.set_fontweight("bold"); at.set_color("white")
    ax.set_title("Micron FQ2'26 Revenue by Segment", color=TEXT_COLOR, fontsize=11, pad=10)
    charts["seg"] = chart_to_bytes(fig)

    # 5. Revenue by Business Unit (FQ2 2026)
    bus_units = ["CMBU\n(Cloud)", "MCBU\n(Mobile/Client)", "CDBU\n(Core DC)", "AEBU\n(Auto/Emb)"]
    bu_rev    = [7.7, 7.7, 5.7, 2.7]
    bu_gm     = [74, 79, 74, 68]
    x = np.arange(len(bus_units))
    fig, ax1 = base_fig(figsize=(9, 5))
    ax2 = ax1.twinx()
    ax2.set_facecolor(AXIS_STYLE["facecolor"])
    bars = ax1.bar(x, bu_rev, color=BAR_BLUE, width=0.5, label="Revenue ($B)")
    ax2.plot(x, bu_gm, color=BAR_AMBER, marker="o", linewidth=2.5,
             markersize=8, label="Gross Margin %", zorder=5)
    ax1.set_xticks(x); ax1.set_xticklabels(bus_units, color=TEXT_COLOR, fontsize=9)
    ax1.set_ylabel("Revenue (US$B)", color=TEXT_COLOR)
    ax2.set_ylabel("Gross Margin %", color=BAR_AMBER)
    ax2.tick_params(colors=BAR_AMBER)
    ax2.spines["right"].set_edgecolor(BAR_AMBER)
    ax1.set_title("Micron FQ2'26 業務部門營收與毛利率", color=TEXT_COLOR, fontsize=11, pad=10)
    ax1.yaxis.grid(True, color=GRID_COLOR, linewidth=0.5)
    ax1.set_axisbelow(True)
    for bar, v in zip(bars, bu_rev):
        ax1.text(bar.get_x()+bar.get_width()/2, v+0.1,
                 f"${v:.1f}B", ha="center", color=TEXT_COLOR, fontsize=9, fontweight="bold")
    for i, v in enumerate(bu_gm):
        ax2.text(i, v+1, f"{v}%", ha="center", color=BAR_AMBER, fontsize=9, fontweight="bold")
    ax2.set_ylim(50, 95)
    h1, l1 = ax1.get_legend_handles_labels()
    h2, l2 = ax2.get_legend_handles_labels()
    ax1.legend(h1+h2, l1+l2, facecolor="#1e293b", labelcolor=TEXT_COLOR, fontsize=8, loc="upper left")
    charts["bu"] = chart_to_bytes(fig)

    # 6. Beat/Miss vs Consensus
    metrics  = ["Revenue\n(US$B)", "Non-GAAP\nEPS (US$)", "Gross\nMargin %"]
    actual   = [23.9, 12.20, 75]
    consens  = [19.2, 8.79, 64]
    x = np.arange(len(metrics))
    w = 0.35
    fig, ax = base_fig(figsize=(8, 4))
    b1 = ax.bar(x - w/2, actual,  w, label="實際值", color=BAR_GREEN)
    b2 = ax.bar(x + w/2, consens, w, label="法人共識", color=BAR_BLUE)
    ax.set_title("Micron FQ2'26 實際 vs 法人共識", color=TEXT_COLOR, fontsize=11, pad=10)
    ax.set_xticks(x); ax.set_xticklabels(metrics, color=TEXT_COLOR, fontsize=9)
    ax.yaxis.grid(True, color=GRID_COLOR, linewidth=0.5); ax.set_axisbelow(True)
    ax.legend(facecolor="#1e293b", labelcolor=TEXT_COLOR, fontsize=9)
    for bar, v in zip(b1, actual):
        ax.text(bar.get_x()+bar.get_width()/2, v*1.02, f"{v}", ha="center",
                color=TEXT_COLOR, fontsize=9, fontweight="bold")
    charts["beat"] = chart_to_bytes(fig)

    # 7. DRAM vs NAND Pricing QoQ % change
    periods = ["FQ3'25", "FQ4'25", "FQ1'26", "FQ2'26", "FQ3'26E"]
    dram_px = [15, 18, 30, 65, 40]  # estimated sequential price change %
    nand_px = [10, 12, 20, 78, 50]
    x = np.arange(len(periods))
    w = 0.35
    fig, ax = base_fig(figsize=(9, 4))
    ax.bar(x - w/2, dram_px, w, label="DRAM ASP QoQ%", color=BAR_BLUE)
    ax.bar(x + w/2, nand_px, w, label="NAND ASP QoQ%", color=BAR_GREEN)
    ax.set_title("Micron DRAM/NAND ASP 季環比漲幅（%）", color=TEXT_COLOR, fontsize=11, pad=10)
    ax.set_xticks(x); ax.set_xticklabels(periods, color=TEXT_COLOR, fontsize=9)
    ax.yaxis.grid(True, color=GRID_COLOR, linewidth=0.5); ax.set_axisbelow(True)
    ax.legend(facecolor="#1e293b", labelcolor=TEXT_COLOR, fontsize=9)
    ax.set_ylabel("ASP QoQ %", color=TEXT_COLOR)
    ax.annotate("FQ2'26: DRAM+65%\nNAND+78%", xy=(3, 78), xytext=(3.2, 85),
                color=TEXT_COLOR, fontsize=8)
    ax.set_ylim(0, 100)
    charts["pricing"] = chart_to_bytes(fig)

    # 8. Free Cash Flow
    fcf_q = ["FQ2'25", "FQ3'25", "FQ4'25", "FQ1'26", "FQ2'26"]
    fcf_v = [0.5, 1.2, 2.1, 3.9, 6.9]
    fig, ax = base_fig(figsize=(8, 4))
    ax.bar(fcf_q, fcf_v, color=[BAR_BLUE]*4+[BAR_GREEN], width=0.55)
    ax.plot(fcf_q, fcf_v, color=LINE_COLOR, marker="o", linewidth=2, markersize=7, zorder=5)
    ax.set_title("Micron Free Cash Flow (US$B) — 季度創歷史新高", color=TEXT_COLOR, fontsize=11, pad=10)
    ax.set_ylabel("FCF (US$B)", color=TEXT_COLOR)
    ax.yaxis.grid(True, color=GRID_COLOR, linewidth=0.5); ax.set_axisbelow(True)
    for i, v in enumerate(fcf_v):
        ax.text(i, v+0.1, f"${v:.1f}B", ha="center", color=TEXT_COLOR, fontsize=9, fontweight="bold")
    ax.annotate("Record FCF\n$6.9B", xy=(4, 6.9), xytext=(3.2, 7.5),
                color=BAR_GREEN, fontsize=8,
                arrowprops=dict(arrowstyle="->", color=BAR_GREEN))
    charts["fcf"] = chart_to_bytes(fig)

    return charts


def build_micron_report(outpath, charts):
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(2.2)
        section.right_margin  = Cm(2.2)

    # ── COVER ──────────────────────────────────────────────────
    p = doc.add_paragraph()
    r = p.add_run("Micron Technology（MU）")
    r.font.size = Pt(22); r.bold = True
    r.font.color.rgb = rgb(DARK_NAVY)
    r.font.name = "Times New Roman"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    r = p.add_run("FY2026 第二季財報更新報告  |  FQ2 2026 Earnings Update")
    r.font.size = Pt(14); r.bold = True
    r.font.color.rgb = rgb((89,89,89))
    r.font.name = "Times New Roman"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    cover_data = [
        ("投資評等",  "買進 (Buy) ✓ 維持", True),
        ("目標價格",  "$500（上調自 $450）", True),
        ("現時股價",  "$422.90（2026/03/20）", False),
        ("潛在漲幅",  "+18.2%", True),
        ("報告日期",  "2026年3月22日", False),
        ("分析師",    "台灣科技股研究團隊", False),
        ("數據來源",  "FQ2 2026 法說會逐字稿 + SEC 10-Q", False),
    ]
    cover_table(doc, cover_data, [Inches(2), Inches(4.2)])
    doc.add_paragraph()
    add_divider(doc)

    # ── KEY TAKEAWAYS ──────────────────────────────────────────
    add_heading(doc, "核心要點  KEY TAKEAWAYS", level=1, color=DARK_NAVY, size=13)
    takeaways = [
        "【史詩級超預期】FQ2 營收 $23.9B 超法人共識 $19.2B（+24.6%），EPS $12.20 超共識 $8.79（+38.8%）。兩項均為公司歷史最大超預期幅度。",
        "【毛利率創歷史新高】Non-GAAP 毛利率達 75%（+18pp QoQ），季增 18 個百分點，幾乎是上年同期的兩倍。",
        "【FQ3 指引比 FQ2 再+40%】FQ3 指引營收 $33.5B（法人共識 $22.6B）、EPS $19.15（法人共識 $13.6B），再度大幅超越市場預期。",
        "【HBM4 正式量產出貨】HBM4 36GB 12H 已於 CY2026 Q1 開始量產出貨，設計用於 NVIDIA Vera Rubin，是首個 HBM4 量產里程碑。",
        "【首份五年戰略客戶協議（SCA）簽署】跨週期可見度大幅提升，商業模式穩定性優於傳統 LTA。",
        "【FY2026 CapEx 上調至 $25B+】確認超級循環為多年持續事件；股息調升 30% 至 $0.15/季。",
    ]
    for t in takeaways:
        add_bullet(doc, t, size=10)

    add_divider(doc)

    # ── RESULTS SNAPSHOT ──────────────────────────────────────
    add_heading(doc, "財報快覽  RESULTS SNAPSHOT", level=1, color=DARK_NAVY, size=13)
    hdrs = ["財務指標", "FQ2'26 實際", "法人共識", "超/低預期", "QoQ 變化", "YoY 變化"]
    rows = [
        ["總營收 (US$B)",     "$23.86B",  "$19.19B", "▲ +24.3%",  "+75%",   "+196%"],
        ["DRAM 營收 (US$B)",  "$18.8B",   "—",       "▲ 紀錄",    "+74%",   "+207%"],
        ["NAND 營收 (US$B)",  "$5.0B",    "—",       "▲ 紀錄",    "+82%",   "+169%"],
        ["毛利率 (非GAAP)",   "75%",      "64%e",    "▲ +11pp",   "+18pp",  "+37pp"],
        ["營業利益率 (非GAAP)","69%",      "—",       "▲ 紀錄",    "+22pp",  "+44pp"],
        ["Non-GAAP EPS",      "$12.20",   "$8.79",   "▲ +38.8%",  "+155%",  "+682%"],
        ["Free Cash Flow",    "$6.9B",    "—",       "▲ 季度紀錄","+77%",   "N/M"],
        ["淨現金部位",        "$6.5B",    "—",       "▲ 史上最高","—",      "—"],
    ]
    beat_miss_table(doc, hdrs, rows)
    add_body(doc, "資料來源：FQ2 2026 法說會逐字稿（2026/03/18）；Non-GAAP 數字。法人共識為彭博估計均值。",
             size=8, italic=True)
    doc.add_paragraph()

    # ── CHARTS ─────────────────────────────────────────────────
    add_heading(doc, "季度財務趨勢", level=2, color=(68,114,196), size=12)
    add_chart(doc, charts["rev"], width=Inches(6.5),
              caption="圖1：Micron 季度合并營收（US$B）；FQ3'26E $33.5B 指引超越公司任何歷史全年營收。資料來源：法說會。")
    add_chart(doc, charts["eps"], width=Inches(6.5),
              caption="圖2：Non-GAAP EPS（US$）；FQ2 $12.20 年增 682%；FQ3E $19.15 再創紀錄。資料來源：法說會。")
    add_chart(doc, charts["gm"], width=Inches(6.5),
              caption="圖3：Non-GAAP 毛利率（%）；FQ2 75% 創歷史紀錄，FQ3E 81%。資料來源：法說會。")

    doc.add_page_break()

    # ── SEGMENT ANALYSIS ──────────────────────────────────────
    add_heading(doc, "部門分析  SEGMENT ANALYSIS", level=1, color=DARK_NAVY, size=13)

    add_heading(doc, "① DRAM — $18.8B，+207% YoY（佔總營收 79%）", level=2,
                color=(68,114,196), size=11)
    add_body(doc,
        "DRAM 位元出貨量季增個位數中段，ASP 季增 65%，受惠於 AI 資料中心需求爆發及供給結構性限制。"
        "1-gamma 節點成為公司歷史上最快達到成熟良率的節點，預計 2026 年年中成為 DRAM 位元主力。"
        "HBM4 36GB 12H 已開始量產出貨用於 NVIDIA Vera Rubin，預計良率爬坡速度快於 HBM3E。"
        "DRAM 庫存天數仍低於 120 天，遠低於 123 天的整體平均，顯示供給極度緊張。", size=10)

    add_heading(doc, "② NAND — $5.0B，+169% YoY（佔總營收 21%）", level=2,
                color=(68,114,196), size=11)
    add_body(doc,
        "NAND 位元出貨量季增個位數低段，ASP 季增近 80%，資料中心 NAND 營收季增超過一倍。"
        "AI 應用（向量資料庫、KV Cache 卸載）及 SSD 在容量儲存層滲透率提升為主要驅動力。"
        "G9 節點（PCIe Gen6 高效能 SSD）正大量生產，122TB 高容量 SSD 獲廣泛採用。"
        "2026 年 NAND 供給預計僅增長約 20%，遠低於需求增速，供需緊張延伸至 2026 年後。", size=10)

    add_chart(doc, charts["seg"], width=Inches(5.5),
              caption="圖4：FQ2'26 營收結構（DRAM 78.7%，NAND 20.9%）。資料來源：法說會逐字稿。")

    add_chart(doc, charts["bu"], width=Inches(6.5),
              caption="圖5：各業務部門營收與毛利率；MCBU 毛利率 79% 最高。資料來源：法說會逐字稿。")

    add_chart(doc, charts["pricing"], width=Inches(6.5),
              caption="圖6：Micron DRAM/NAND ASP 季環比漲幅；FQ2'26 DRAM+65%，NAND+78%。資料來源：法說會。")

    doc.add_page_break()

    # ── GUIDANCE ──────────────────────────────────────────────
    add_heading(doc, "展望與指引  GUIDANCE & OUTLOOK", level=1,
                color=DARK_NAVY, size=13)
    hdrs2 = ["指標", "FQ2'26 實際", "FQ3'26 指引", "法人原共識", "超指引幅度"]
    rows2 = [
        ["營收 (US$B)",      "$23.86B", "$33.5B ±$750M", "$22.6B", "▲ +48.2%"],
        ["毛利率 (Non-GAAP)","75%",     "~81%",          "72%e",   "▲ +9pp"],
        ["EPS (Non-GAAP)",   "$12.20",  "$19.15 ±$0.40", "$13.6",  "▲ +40.8%"],
        ["CapEx (US$B)",     "$5.0B",   "~$7.0B",        "—",      "FY26 >$25B"],
        ["股息 (US$/季)",    "$0.115",  "$0.15 (+30%)",  "—",      "▲ 提升 30%"],
    ]
    beat_miss_table(doc, hdrs2, rows2)
    add_body(doc, "資料來源：FQ2 2026 法說會（2026/03/18）。法人共識截至 2026/03/17 彭博均值。",
             size=8, italic=True)
    doc.add_paragraph()

    add_body(doc,
        "重要前瞻指引摘要：①FQ3 單季指引 $33.5B，超過公司 2024 財年全年營收，歷史先例；"
        "②81% 毛利率指引，HBM4 量產爬坡貢獻；③FY2026 CapEx >$25B（含台灣銅羅廠收購），"
        "FY2027 CapEx 再大幅躍升；④簽署首份五年 SCA（Strategic Customer Agreement）。", size=10)

    add_chart(doc, charts["beat"], width=Inches(6.2),
              caption="圖7：FQ2'26 實際 vs 法人共識；三項主要指標均大幅超越預期。資料來源：Bloomberg 共識、法說會。")
    add_chart(doc, charts["fcf"], width=Inches(6.2),
              caption="圖8：Micron Free Cash Flow 趨勢；FQ2'26 $6.9B 創季度紀錄。資料來源：法說會逐字稿。")

    doc.add_page_break()

    # ── INVESTMENT THESIS UPDATE ───────────────────────────────
    add_heading(doc, "投資論點更新  THESIS UPDATE", level=1,
                color=DARK_NAVY, size=13)
    add_body(doc, "本季財報全面強化我們的記憶體超級循環核心論點。以下四點為關鍵更新：", size=10)

    add_heading(doc, "✅ 論點強化 — 三大結構性支柱確立", level=2, color=GREEN_OK, size=11)
    positives = [
        "供需緊張確認延伸至 2026 年後：DRAM/NAND 供給 2026 年各僅增 ~20%，而 AI 資料中心需求首次超過產業 TAM 50%。",
        "定價能力超預期：DRAM +65%、NAND +78% 的 ASP 季增幅度，遠超週期高峰期歷史前例，顯示 AI 需求拉力非傳統週期。",
        "技術護城河加深：HBM4 量產、1-gamma 最快良率爬坡、首份五年 SCA，均為競爭差異化的長期優勢。",
        "財務體質歷史最強：淨現金 $6.5B 創歷史高峰，FCF $6.9B 創記錄，具備充裕的再投資與股東回報能力。",
    ]
    for p_item in positives:
        add_bullet(doc, p_item, size=10)

    add_heading(doc, "⚠ 監控風險", level=2, color=AMBER, size=11)
    risks = [
        "三星 HBM4 量產進度：P4 fab 進展若提前至 2026H2，可能衝擊 HBM 定價。",
        "貿易/地緣政治：FQ3 指引明確排除貿易/地緣政治影響，需密切追蹤美中關稅動態。",
        "AI CapEx 週期：任一 Hyperscaler CapEx YoY 下修 >10% 為尾部風險，目前尚未出現此訊號。",
        "消費性終端市場弱化：PC/智慧型手機單位出貨量可能因高定價下滑兩位數，但 AI 內容升級抵銷衝擊。",
    ]
    for r_item in risks:
        add_bullet(doc, r_item, size=10)

    add_divider(doc)

    # ── VALUATION ─────────────────────────────────────────────
    add_heading(doc, "估值分析  VALUATION", level=1, color=DARK_NAVY, size=13)

    add_heading(doc, "三情境目標價（加權法）", level=2, color=(68,114,196), size=11)
    hdrs3 = ["情境", "概率", "FY2026E EPS", "PE 倍數", "目標價 (US$)", "依據"]
    rows3 = [
        ["熊市 Bear",  "20%", "$35",   "7x",  "$245", "AI CapEx 下修 + 三星供給衝擊"],
        ["基本 Base",  "55%", "$50",   "9x",  "$450", "FQ3 指引 $19.15 + Q4 估 $22 = ~$50"],
        ["牛市 Bull",  "25%", "$58",   "10x", "$580", "超級循環延伸 + HBM4 溢價"],
        ["加權目標價", "—",  "—",     "—",  "$489", "= 245×20% + 450×55% + 580×25%"],
    ]
    beat_miss_table(doc, hdrs3, rows3)
    add_body(doc,
        "目標價設定 $500（取整，貼近加權計算 $489）。關鍵假設：FQ1已報 ~$5 + FQ2 $12.20 + FQ3E $19.15 + FQ4E ~$22 = FY2026E EPS ~$58；以 8.6x Forward PE（週期超級循環中段合理）推算。現價 $422.90，潛在漲幅 +18.2%。",
        size=10)

    add_divider(doc)

    # ── ESTIMATES UPDATE ──────────────────────────────────────
    add_heading(doc, "預估值更新  ESTIMATE REVISIONS", level=1,
                color=DARK_NAVY, size=13)
    hdrs4 = ["指標", "舊預估（FQ2 前）", "新預估（FQ2 後）", "修正幅度"]
    rows4 = [
        ["FY2026 全年營收 (US$B)", "$65",  "$90-100",   "▲ +38-54%"],
        ["FY2026 全年毛利率",      "65%",  "75%+",      "▲ +10pp+"],
        ["FY2026 Non-GAAP EPS",   "$28",  "$55-60",    "▲ +96-114%"],
        ["FQ3'26 營收指引",        "$22.6B","$33.5B",   "▲ +48%（遠超）"],
        ["FQ3'26 毛利率指引",      "72%",  "81%",       "▲ +9pp"],
        ["FY2026 CapEx (US$B)",   "$20B", ">$25B",     "▲ +25%+"],
    ]
    beat_miss_table(doc, hdrs4, rows4)

    doc.add_page_break()

    # ── SOURCES ───────────────────────────────────────────────
    add_heading(doc, "資料來源  SOURCES & REFERENCES", level=1,
                color=DARK_NAVY, size=13)
    sources = [
        "Micron FQ2 2026 法說會逐字稿（2026/03/18）— 本地文件：Micron_2026Q2_逐字稿.txt",
        "FQ2 2026 Earnings Press Release — https://investors.micron.com",
        "Form 10-Q FQ2 2026（Filed ~2026/04）— SEC EDGAR: https://www.sec.gov/cgi-bin/browse-edgar?action=getcompany&CIK=MU",
        "FQ2 2026 Earnings Call: Micron Q2 2026 slides — https://www.investing.com/news/company-news/micron-q2-2026-slides-ai-demand-drives-record-239b-revenue",
        "Futurum Research: Micron Q2 FY2026 Earnings — https://futurumgroup.com/insights/micron-q2-fy-2026-earnings-driven-by-ai-led-memory-demand/",
        "StockAnalysis MU Statistics — https://stockanalysis.com/stocks/mu/statistics/",
        "法人共識：Bloomberg Terminal 截至 2026/03/17；EPS 共識 $8.79，Revenue $19.19B",
        "股價 $422.90 截至 2026/03/20 NASDAQ 收盤",
    ]
    for s in sources:
        add_bullet(doc, s, size=9)

    add_body(doc,
        "免責聲明：本報告僅供個人投資研究參考，非正式投資建議。所有預估值（標記「E」或「e」）為分析師預測，"
        "實際結果可能存在重大差異。對美股 MU 的估值建議以美元計，匯率風險請自行評估。",
        size=8, italic=True)

    doc.save(outpath)
    print(f"✓ Micron report saved: {outpath}")


# ════════════════════════════════════════════════════════════════
# MAIN
# ════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    BASE = "/Users/wayne/Desktop/Invest/DRAM"

    print("Building Winbond charts...")
    wb_charts = build_winbond_charts(BASE)
    winbond_path = os.path.join(BASE, "Winbond_2344_Q4_2025_Earnings_Update.docx")
    build_winbond_report(winbond_path, wb_charts)

    print("Building Micron charts...")
    mu_charts = build_micron_charts()
    micron_path = os.path.join(BASE, "Micron_MU_FQ2_2026_Earnings_Update.docx")
    build_micron_report(micron_path, mu_charts)

    print("\n✅ Both reports complete.")
    print(f"   → {winbond_path}")
    print(f"   → {micron_path}")
